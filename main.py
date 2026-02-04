# main.py
# AAP išdavimas – viename faile (stabilus controller, be Tk klaidų)
#
# Failai (pagal settings.json):
# - Darbuotojai.xlsx
# - AAP DB.xlsx
# - Aprangos kodai.xlsx
# - Word šablonas (template)
# - settings.json
#
# Pastabos:
# - Naujos įrangos lentelė: Kodas | Susidėvėjimas | Apranga (NEredaguojama)
# - Kodas su dydžiu: 05-32 -> Excel kodas rašomas "05", o pavadinime pridedama "(32 dydis)"
# - Jei kodas nerandamas "Aprangos kodai.xlsx" – iššoka mažas dialogas pavadinimui įvesti (lentelėje pavadinimas lieka neredaguojamas)
# - Tab eiga: Kodas -> Susidėvėjimas -> kitos eilutės Kodas

import json
import os
from dataclasses import dataclass
from datetime import date, datetime, timedelta
from pathlib import Path
import tkinter as tk
from tkinter import ttk, messagebox, simpledialog

from openpyxl import Workbook, load_workbook


# -----------------------------
# Excel stulpeliai
# -----------------------------
AAP_DB_COLUMNS = [
    "Numeris",
    "Vardas Pavardė",
    "Tab. Nr",
    "Padalinys",
    "Pareigos",
    "Lytis",
    "Aprangos kodas",
    "Apranga",
    "Išduota",
    "Susidėvėjimas",
]

DARBUOTOJAI_COLUMNS = ["Vardas", "Pavardė", "TabNr", "Pareigos", "Padalinys", "Lytis"]

GEAR_EXCEL_HEADERS = ["Prekės Nr.", "Prekės pavadinimas"]


# -----------------------------
# Modeliai
# -----------------------------
@dataclass
class EmployeeInfo:
    tab_nr: str
    name: str
    department: str
    position: str
    gender: str
    issue_date: date | None = None


# -----------------------------
# Pagalbinės
# -----------------------------
def parse_date(value):
    if value is None:
        return None
    if isinstance(value, datetime):
        return value.date()
    if isinstance(value, date):
        return value
    s = str(value).strip()
    for fmt in ("%Y-%m-%d", "%d.%m.%Y", "%Y/%m/%d", "%Y.%m.%d"):
        try:
            return datetime.strptime(s, fmt).date()
        except ValueError:
            pass
    return None


def add_months(start_date: date, months: int) -> date | None:
    if start_date is None:
        return None
    month = start_date.month - 1 + months
    year = start_date.year + month // 12
    month = month % 12 + 1

    if month == 12:
        first_next = date(year + 1, 1, 1)
    else:
        first_next = date(year, month + 1, 1)
    last_day = (first_next - timedelta(days=1)).day

    day = min(start_date.day, last_day)
    return date(year, month, day)


def resolve_path(base_path: Path, target: str) -> Path:
    p = Path(target)
    return p if p.is_absolute() else (base_path / p).resolve()


def split_code_and_size(code: str):
    c = (code or "").strip()
    if "-" in c:
        base, size = c.split("-", 1)
        return base.strip(), size.strip()
    return c, ""


def strip_size_suffix(name: str) -> str:
    if not name:
        return ""
    t = name.strip()
    if "(" in t and t.endswith(")"):
        return t[: t.rfind("(")].strip()
    return t


def ensure_size_suffix(name: str, size: str) -> str:
    if not size:
        return name
    base = strip_size_suffix(name)
    return f"{base} ({size} dydis)".strip()


def replace_placeholders(doc, mapping: dict[str, str]):
    for p in doc.paragraphs:
        for k, v in mapping.items():
            if k in p.text:
                for run in p.runs:
                    run.text = run.text.replace(k, v)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for k, v in mapping.items():
                        if k in p.text:
                            for run in p.runs:
                                run.text = run.text.replace(k, v)


# -----------------------------
# Settings (JSON)
# -----------------------------
def load_settings(settings_path: str | Path):
    settings_path = Path(settings_path)
    with open(settings_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    base_dir = settings_path.resolve().parent
    paths = data.get("paths", {})

    return {
        "template": resolve_path(base_dir, paths.get("template", "template.docx")),
        "aap_db": resolve_path(base_dir, paths.get("aap_db", paths.get("excel", "AAP DB.xlsx"))),
        "darbuotojai": resolve_path(base_dir, paths.get("darbuotojai", paths.get("employees", "Darbuotojai.xlsx"))),
        "gear_excel": resolve_path(base_dir, paths.get("gear_excel", "Aprangos kodai.xlsx")),
        "outputs": resolve_path(base_dir, paths.get("outputs", "sugeneruotos kortelės")),
    }


# -----------------------------
# Excel užtikrinimai
# -----------------------------
def ensure_excel_aap_db(path: Path):
    if path.exists():
        return
    wb = Workbook()
    ws = wb.active
    ws.append(AAP_DB_COLUMNS)
    wb.save(path)


def ensure_excel_darbuotojai(path: Path):
    if path.exists():
        return
    wb = Workbook()
    ws = wb.active
    ws.append(DARBUOTOJAI_COLUMNS)
    wb.save(path)


def ensure_gear_excel_file(path: Path):
    if path.exists():
        return
    wb = Workbook()
    ws = wb.active
    ws.title = "Aprangos kodai"
    ws.append(GEAR_EXCEL_HEADERS)
    wb.save(path)


# -----------------------------
# Gear Excel repo
# -----------------------------
def load_gear_codes_from_excel(path: Path) -> dict:
    """
    { "ausi101": {"name":"Ausinės"}, ... }
    """
    if not path.exists():
        return {}
    wb = load_workbook(path)
    ws = wb.active

    headers = [str(c.value).strip() if c.value is not None else "" for c in ws[1]]
    code_idx = headers.index("Prekės Nr.") + 1 if "Prekės Nr." in headers else 1
    name_idx = headers.index("Prekės pavadinimas") + 1 if "Prekės pavadinimas" in headers else 2

    out = {}
    for r in range(2, ws.max_row + 1):
        code = ws.cell(r, code_idx).value
        name = ws.cell(r, name_idx).value
        if code is None or name is None:
            continue
        code = str(code).strip()
        name = str(name).strip()
        if code and name:
            out[code] = {"name": name}
    return out


def add_gear_code_to_excel(path: Path, code: str, name: str):
    code = str(code or "").strip()
    name = str(name or "").strip()
    if not code or not name:
        return

    ensure_gear_excel_file(path)
    wb = load_workbook(path)
    ws = wb.active

    existing = set()
    for r in range(2, ws.max_row + 1):
        v = ws.cell(r, 1).value
        if v is not None:
            existing.add(str(v).strip())

    if code in existing:
        return

    ws.append([code, name])
    wb.save(path)


# -----------------------------
# Darbuotojai repo
# -----------------------------
class DarbuotojaiRepo:
    def __init__(self, path: Path):
        self.path = path
        ensure_excel_darbuotojai(self.path)

    def _open(self):
        wb = load_workbook(self.path)
        return wb, wb.active

    def load_all(self) -> list[dict]:
        wb, ws = self._open()
        out = []
        for row in ws.iter_rows(min_row=2, values_only=True):
            if not any(row):
                continue
            vardas = str(row[0] or "").strip()
            pavarde = str(row[1] or "").strip()
            tab = str(row[2] or "").strip()
            pareigos = str(row[3] or "").strip()
            padalinys = str(row[4] or "").strip()
            lytis = str(row[5] or "").strip()
            out.append({
                "Vardas": vardas,
                "Pavardė": pavarde,
                "TabNr": tab,
                "Pareigos": pareigos,
                "Padalinys": padalinys,
                "Lytis": lytis,
            })
        return out

    def list_departments(self) -> list[str]:
        rows = self.load_all()
        return sorted({r["Padalinys"] for r in rows if r["Padalinys"]})

    def list_positions_for_department(self, dept: str) -> list[str]:
        dept = (dept or "").strip()
        rows = self.load_all()
        return sorted({r["Pareigos"] for r in rows if r["Padalinys"] == dept and r["Pareigos"]})

    def find_by_tab(self, tab_nr: str) -> dict | None:
        tab_nr = str(tab_nr or "").strip()
        if not tab_nr:
            return None
        for r in self.load_all():
            if str(r.get("TabNr", "")).strip() == tab_nr:
                return r
        return None

    def upsert(self, e: EmployeeInfo):
        tab_nr = str(e.tab_nr).strip()
        vp = str(e.name).strip()
        if not tab_nr or not vp:
            return

        parts = vp.split()
        vardas = parts[0] if parts else vp
        pavarde = " ".join(parts[1:]) if len(parts) > 1 else ""

        wb, ws = self._open()

        found = None
        for i in range(2, ws.max_row + 1):
            tab = str(ws.cell(i, 3).value or "").strip()
            if tab == tab_nr:
                found = i
                break

        if found is None:
            ws.insert_rows(2)
            r = 2
        else:
            r = found

        ws.cell(r, 1, vardas)
        ws.cell(r, 2, pavarde)
        ws.cell(r, 3, tab_nr)
        ws.cell(r, 4, e.position)
        ws.cell(r, 5, e.department)
        ws.cell(r, 6, e.gender)

        wb.save(self.path)


# -----------------------------
# AAP DB repo
# -----------------------------
class AAPDbRepo:
    def __init__(self, path: Path):
        self.path = path
        ensure_excel_aap_db(self.path)

    def _open(self):
        wb = load_workbook(self.path)
        return wb, wb.active

    def load_rows(self) -> list[dict]:
        wb, ws = self._open()
        out = []
        for row in ws.iter_rows(min_row=2, values_only=True):
            if not any(row):
                continue
            out.append(dict(zip(AAP_DB_COLUMNS, row)))
        return out

    def next_document_number(self) -> int:
        rows = self.load_rows()
        nums = []
        for r in rows:
            v = r.get("Numeris")
            try:
                if v not in (None, ""):
                    nums.append(int(v))
            except Exception:
                pass
        return (max(nums) if nums else 0) + 1

    def append_issuance(self, employee: EmployeeInfo, issued_items: list[dict], doc_no: int, issued_date: date):
        wb, ws = self._open()

        for it in reversed(issued_items):
            ws.insert_rows(2)
            ws.cell(2, 1, doc_no)
            ws.cell(2, 2, employee.name)
            ws.cell(2, 3, employee.tab_nr)
            ws.cell(2, 4, employee.department)
            ws.cell(2, 5, employee.position)
            ws.cell(2, 6, employee.gender)
            ws.cell(2, 7, it["code"])  # be "-dydis"
            ws.cell(2, 8, it["name"])
            ws.cell(2, 9, issued_date.isoformat())
            ws.cell(2, 10, int(it.get("months", 0) or 0))

        wb.save(self.path)

    def current_gear_for_employee(self, employee: EmployeeInfo) -> list[dict]:
        rows = self.load_rows()
        rel = [r for r in rows if str(r.get("Tab. Nr", "")).strip() == employee.tab_nr]

        # tik dabartinis padalinys/pareigos
        rel = [r for r in rel
               if str(r.get("Padalinys", "")).strip() == employee.department
               and str(r.get("Pareigos", "")).strip() == employee.position]

        latest = {}
        for r in rel:
            nm = str(r.get("Apranga", "") or "").strip()
            if not nm:
                continue
            base = strip_size_suffix(nm)
            issued = parse_date(r.get("Išduota"))
            if not issued:
                continue
            prev = latest.get(base)
            if prev is None or issued > prev["issued"]:
                latest[base] = {"row": r, "issued": issued}

        today = date.today()
        out = []
        for base, info in latest.items():
            r = info["row"]
            issued = parse_date(r.get("Išduota"))
            months = int(r.get("Susidėvėjimas") or 0)
            change_date = add_months(issued, months) if issued else None
            remaining = (change_date - today).days if change_date else None
            out.append({
                "Apranga": str(r.get("Apranga", "") or ""),
                "Kodas": str(r.get("Aprangos kodas", "") or ""),
                "Išduota": issued,
                "Mėn": months,
                "Keisti iki": change_date,
                "Likę": remaining,
            })

        out.sort(key=lambda x: (x["Likę"] is None, x["Likę"] if x["Likę"] is not None else 10**9))
        return out


# -----------------------------
# Popup: sąrašas (platesnis)
# -----------------------------
class ListSelectPopup(tk.Toplevel):
    def __init__(self, parent, title: str, values: list[str], on_pick):
        super().__init__(parent)
        self.title(title)
        self.resizable(True, True)
        self.minsize(620, 460)
        self.geometry("720x520")
        self.attributes("-topmost", True)
        self.on_pick = on_pick

        frm = ttk.Frame(self)
        frm.pack(fill="both", expand=True, padx=10, pady=10)

        self.listbox = tk.Listbox(frm, height=18, width=85, exportselection=False)
        scroll = ttk.Scrollbar(frm, orient="vertical", command=self.listbox.yview)
        self.listbox.configure(yscrollcommand=scroll.set)

        self.listbox.grid(row=0, column=0, sticky="nsew")
        scroll.grid(row=0, column=1, sticky="ns")

        frm.rowconfigure(0, weight=1)
        frm.columnconfigure(0, weight=1)

        for v in values:
            self.listbox.insert(tk.END, v)

        self.listbox.bind("<Double-Button-1>", self._choose)
        self.listbox.bind("<Return>", self._choose)

        btns = ttk.Frame(self)
        btns.pack(pady=(0, 10))
        ttk.Button(btns, text="Pasirinkti", command=self._choose).pack(side="left", padx=6)
        ttk.Button(btns, text="Uždaryti", command=self.destroy).pack(side="left", padx=6)

    def _choose(self, _ev=None):
        sel = self.listbox.curselection()
        if not sel:
            return
        v = self.listbox.get(sel[0])
        self.on_pick(v)
        self.destroy()


# -----------------------------
# Popup: kalendorius (paprastas ir stabilus)
# -----------------------------
class CalendarPopup(tk.Toplevel):
    def __init__(self, parent, initial_date: date, on_pick):
        super().__init__(parent)
        self.on_pick = on_pick
        self.cur = date(initial_date.year, initial_date.month, 1)
        self.title("Pasirinkti datą")
        self.resizable(False, False)
        self.attributes("-topmost", True)

        top = ttk.Frame(self)
        top.pack(padx=10, pady=8, fill="x")

        ttk.Button(top, text="◀", width=3, command=self.prev_month).pack(side="left")
        self.lbl = ttk.Label(top, text="", width=22, anchor="center")
        self.lbl.pack(side="left", padx=8)
        ttk.Button(top, text="▶", width=3, command=self.next_month).pack(side="left")

        self.grid_frame = ttk.Frame(self)
        self.grid_frame.pack(padx=10, pady=(0, 10))

        for i, h in enumerate(["Pr", "An", "Tr", "Kt", "Pn", "Št", "Sk"]):
            ttk.Label(self.grid_frame, text=h, width=4, anchor="center").grid(row=0, column=i, padx=1, pady=1)

        self.draw()

    def month_name(self, m):
        names = ["Sausis", "Vasaris", "Kovas", "Balandis", "Gegužė", "Birželis",
                 "Liepa", "Rugpjūtis", "Rugsėjis", "Spalis", "Lapkritis", "Gruodis"]
        return names[m - 1]

    def _days_in_month(self, y, m):
        if m == 12:
            first_next = date(y + 1, 1, 1)
        else:
            first_next = date(y, m + 1, 1)
        return (first_next - timedelta(days=1)).day

    def prev_month(self):
        y, m = self.cur.year, self.cur.month - 1
        if m == 0:
            m = 12
            y -= 1
        self.cur = date(y, m, 1)
        self.draw()

    def next_month(self):
        y, m = self.cur.year, self.cur.month + 1
        if m == 13:
            m = 1
            y += 1
        self.cur = date(y, m, 1)
        self.draw()

    def draw(self):
        for w in list(self.grid_frame.grid_slaves()):
            if int(w.grid_info().get("row", 0)) >= 1:
                w.destroy()

        self.lbl.config(text=f"{self.cur.year} {self.month_name(self.cur.month)}")

        first_wd = self.cur.weekday()  # Mon=0
        last_day = self._days_in_month(self.cur.year, self.cur.month)

        r = 1
        c = first_wd
        for d in range(1, last_day + 1):
            dd = date(self.cur.year, self.cur.month, d)
            b = ttk.Button(self.grid_frame, text=f"{d:02d}", width=4, command=lambda x=dd: self.pick(x))
            b.grid(row=r, column=c, padx=1, pady=1)
            c += 1
            if c >= 7:
                c = 0
                r += 1

    def pick(self, d: date):
        self.on_pick(d)
        self.destroy()


# -----------------------------
# App
# -----------------------------
class AAPApp(tk.Tk):
    def __init__(self, settings_path: str | Path):
        super().__init__()

        self.settings = load_settings(settings_path)

        ensure_excel_darbuotojai(self.settings["darbuotojai"])
        ensure_excel_aap_db(self.settings["aap_db"])
        ensure_gear_excel_file(self.settings["gear_excel"])

        self.repo_people = DarbuotojaiRepo(self.settings["darbuotojai"])
        self.repo_db = AAPDbRepo(self.settings["aap_db"])

        self.gear_codes = load_gear_codes_from_excel(self.settings["gear_excel"])

        self.current_employee: EmployeeInfo | None = None
        self.change_mode = False

        self.title("AAP Issuance")
        self.geometry("900x650")
        self.resizable(False, False)
        self._center()

        container = ttk.Frame(self)
        container.pack(fill="both", expand=True)

        self.pages = {}
        for PageClass in (EmployeeSelectPage, EmployeeInfoPage, CurrentGearPage, NewGearPage):
            page = PageClass(container, self)  # <-- Svarbiausia: controller yra self (AAPApp)
            self.pages[PageClass.__name__] = page
            page.grid(row=0, column=0, sticky="nsew")

        self.show("EmployeeSelectPage")

    def _center(self):
        self.update_idletasks()
        w = self.winfo_width()
        h = self.winfo_height()
        x = (self.winfo_screenwidth() // 2) - (w // 2)
        y = (self.winfo_screenheight() // 2) - (h // 2)
        self.geometry(f"{w}x{h}+{x}+{y}")

    def show(self, page_name: str):
        page = self.pages[page_name]
        if hasattr(page, "refresh"):
            page.refresh()
        page.tkraise()

    # --- People helpers ---
    def list_employees(self) -> list[EmployeeInfo]:
        rows = self.repo_people.load_all()
        out = []
        for r in rows:
            tab = str(r.get("TabNr", "")).strip()
            if not tab:
                continue
            name = f"{r.get('Vardas','')}".strip() + " " + f"{r.get('Pavardė','')}".strip()
            name = " ".join(name.split()).strip()
            out.append(EmployeeInfo(
                tab_nr=tab,
                name=name,
                department=str(r.get("Padalinys", "") or "").strip(),
                position=str(r.get("Pareigos", "") or "").strip(),
                gender=str(r.get("Lytis", "") or "").strip() or "Vyras",
            ))
        out.sort(key=lambda e: e.tab_nr)
        return out

    def find_employee_by_tab(self, tab_nr: str) -> EmployeeInfo | None:
        r = self.repo_people.find_by_tab(tab_nr)
        if not r:
            return None
        name = f"{r.get('Vardas','')}".strip() + " " + f"{r.get('Pavardė','')}".strip()
        name = " ".join(name.split()).strip()
        return EmployeeInfo(
            tab_nr=str(r.get("TabNr", "")).strip(),
            name=name,
            department=str(r.get("Padalinys", "") or "").strip(),
            position=str(r.get("Pareigos", "") or "").strip(),
            gender=str(r.get("Lytis", "") or "").strip() or "Vyras",
        )

    def upsert_employee(self, e: EmployeeInfo):
        self.repo_people.upsert(e)

    def list_departments(self) -> list[str]:
        return self.repo_people.list_departments()

    def list_positions_for_department(self, dept: str) -> list[str]:
        return self.repo_people.list_positions_for_department(dept)

    # --- Gear helpers ---
    def reload_gear_codes(self):
        self.gear_codes = load_gear_codes_from_excel(self.settings["gear_excel"])

    def ensure_gear_code_exists(self, code_base: str, name_base: str):
        code_base = (code_base or "").strip()
        name_base = strip_size_suffix(name_base or "").strip()
        if not code_base or not name_base:
            return
        if code_base not in self.gear_codes:
            add_gear_code_to_excel(self.settings["gear_excel"], code_base, name_base)
            self.reload_gear_codes()

    # --- Word generation ---
    def generate_word_doc(self, employee: EmployeeInfo, issued_items: list[dict], doc_no: int, issued_date: date):
        try:
            from docx import Document
        except Exception:
            messagebox.showerror("Trūksta bibliotekos", "Reikia įdiegti python-docx (paketas: python-docx).")
            return

        template_path = self.settings["template"]
        out_dir = self.settings["outputs"]
        out_dir.mkdir(parents=True, exist_ok=True)

        doc = Document(template_path)

        replace_placeholders(doc, {
            "{Employee}": employee.name,
            "{Emploee}": employee.name,
            "{Departament}": employee.department,
            "{Department}": employee.department,
            "{Position}": employee.position,
        })

        if doc.tables:
            table = doc.tables[0]

            # paliekam tik header eilutę
            while len(table.rows) > 1:
                table._tbl.remove(table.rows[1]._tr)

            for it in issued_items[:14]:
                row = table.add_row()
                cells = row.cells
                n = len(cells)

                months = int(it.get("months", 0) or 0)
                change_date = add_months(issued_date, months)

                # stulpelių pildymas pagal tai, kiek jų yra šablone
                # 0: Išduota
                if n >= 1:
                    cells[0].text = issued_date.strftime("%Y-%m-%d")

                # 1: Kodas
                if n >= 2:
                    cells[1].text = str(it.get("code", "") or "")

                # 2: Apranga
                if n >= 3:
                    cells[2].text = str(it.get("name", "") or "")

                # 3: Susidėvėjimas (mėn.)
                if n >= 4:
                    cells[3].text = str(months) if months else ""

                # 4: Susidėvėjimo data / Keisti iki
                if n >= 5:
                    cells[4].text = change_date.strftime("%Y-%m-%d") if change_date else ""

                # 5: Gavėjo parašas (jei šablone yra toks stulpelis)
                if n >= 6:
                    cells[5].text = "__"

        filename = f"AAP {doc_no} {employee.name}.docx"
        doc.save(out_dir / filename)


# -----------------------------
# Page 1: select employee
# -----------------------------
class EmployeeSelectPage(ttk.Frame):
    def __init__(self, parent, app: AAPApp):
        super().__init__(parent)
        self.app = app

        ttk.Label(self, text="1) Darbuotojo paieška", font=("Arial", 16)).pack(pady=10)

        f = ttk.Frame(self)
        f.pack(fill="x", padx=20)
        ttk.Label(f, text="Paieška (TabNr arba vardas):").pack(side="left")
        self.q = tk.StringVar()
        self.q.trace_add("write", lambda *_: self._fill())
        ttk.Entry(f, textvariable=self.q, width=40).pack(side="left", padx=8)

        lf = ttk.Frame(self)
        lf.pack(fill="both", expand=True, padx=20, pady=10)
        self.listbox = tk.Listbox(lf, height=18)
        self.listbox.pack(side="left", fill="both", expand=True)
        sb = ttk.Scrollbar(lf, orient="vertical", command=self.listbox.yview)
        sb.pack(side="right", fill="y")
        self.listbox.config(yscrollcommand=sb.set)

        bf = ttk.Frame(self)
        bf.pack(pady=10)
        ttk.Button(bf, text="Toliau", command=self.on_next).pack(side="left", padx=6)
        ttk.Button(bf, text="Naujas darbuotojas", command=self.on_new).pack(side="left", padx=6)
        ttk.Button(bf, text="Keisti padalinį", command=self.on_change).pack(side="left", padx=6)

        self.items: list[EmployeeInfo] = []

    def refresh(self):
        self._fill()

    def _fill(self):
        all_emp = self.app.list_employees()
        q = (self.q.get() or "").strip().lower()

        self.items = []
        self.listbox.delete(0, tk.END)

        for e in all_emp:
            label = f"{e.tab_nr} — {e.name}"
            if q and (q not in e.tab_nr.lower() and q not in e.name.lower()):
                continue
            self.items.append(e)
            self.listbox.insert(tk.END, label)

    def _selected(self) -> EmployeeInfo | None:
        sel = self.listbox.curselection()
        return self.items[sel[0]] if sel else None

    def on_next(self):
        e = self._selected()
        if not e:
            messagebox.showwarning("Pasirinkimas", "Pasirink darbuotoją iš sąrašo.")
            return
        self.app.current_employee = e
        self.app.change_mode = False
        self.app.show("CurrentGearPage")

    def on_new(self):
        self.app.current_employee = None
        self.app.change_mode = False
        self.app.show("EmployeeInfoPage")

    def on_change(self):
        e = self._selected()
        if not e:
            messagebox.showwarning("Pasirinkimas", "Pasirink darbuotoją iš sąrašo.")
            return
        self.app.current_employee = e
        self.app.change_mode = True
        self.app.show("EmployeeInfoPage")


# -----------------------------
# Page 2: employee info
# -----------------------------
class EmployeeInfoPage(ttk.Frame):
    def __init__(self, parent, app: AAPApp):
        super().__init__(parent)
        self.app = app

        ttk.Label(self, text="2) Darbuotojo informacija", font=("Arial", 16)).pack(pady=10)

        form = ttk.Frame(self)
        form.pack(pady=10)

        self.tab_var = tk.StringVar()
        self.name_var = tk.StringVar()
        self.gender_var = tk.StringVar(value="Vyras")
        self.dept_var = tk.StringVar()
        self.pos_var = tk.StringVar()

        self.issue_date_var = tk.StringVar(value=date.today().isoformat())
        self.issue_date = date.today()

        ttk.Label(form, text="Tab. Nr:").grid(row=0, column=0, sticky="e", padx=5, pady=6)
        self.tab_entry = ttk.Entry(form, textvariable=self.tab_var, width=52)
        self.tab_entry.grid(row=0, column=1, sticky="w", padx=5, pady=6)

        ttk.Label(form, text="Vardas Pavardė:").grid(row=1, column=0, sticky="e", padx=5, pady=6)
        self.name_entry = ttk.Entry(form, textvariable=self.name_var, width=52)
        self.name_entry.grid(row=1, column=1, sticky="w", padx=5, pady=6)

        ttk.Label(form, text="Lytis:").grid(row=2, column=0, sticky="e", padx=5, pady=6)
        self.gender_combo = ttk.Combobox(form, textvariable=self.gender_var, values=["Vyras", "Moteris"],
                                         state="readonly", width=49)
        self.gender_combo.grid(row=2, column=1, sticky="w", padx=5, pady=6)

        ttk.Label(form, text="Padalinys:").grid(row=3, column=0, sticky="e", padx=5, pady=6)
        dept_row = ttk.Frame(form)
        dept_row.grid(row=3, column=1, sticky="w", padx=5, pady=6)
        self.dept_combo = ttk.Combobox(dept_row, textvariable=self.dept_var, values=[], state="readonly", width=44)
        self.dept_combo.pack(side="left")
        ttk.Button(dept_row, text="Rodyti sąrašą", command=self.open_dept_list).pack(side="left", padx=8)
        self.dept_combo.bind("<<ComboboxSelected>>", lambda _e: self._on_dept_change())

        ttk.Label(form, text="Pareigos:").grid(row=4, column=0, sticky="e", padx=5, pady=6)
        pos_row = ttk.Frame(form)
        pos_row.grid(row=4, column=1, sticky="w", padx=5, pady=6)
        self.pos_combo = ttk.Combobox(pos_row, textvariable=self.pos_var, values=[], state="readonly", width=44)
        self.pos_combo.pack(side="left")
        ttk.Button(pos_row, text="Rodyti sąrašą", command=self.open_pos_list).pack(side="left", padx=8)

        # Data (tik naujam darbuotojui)
        self.date_row = ttk.Frame(form)
        ttk.Label(self.date_row, text="Išdavimo data:").pack(side="left", padx=(0, 6))
        self.date_entry = ttk.Entry(self.date_row, textvariable=self.issue_date_var, width=16, state="readonly")
        self.date_entry.pack(side="left")
        ttk.Button(self.date_row, text="📅", command=self.pick_date).pack(side="left", padx=6)
        ttk.Button(self.date_row, text="Šiandien", command=self.set_today).pack(side="left")

        bf = ttk.Frame(self)
        bf.pack(pady=12)
        ttk.Button(bf, text="Atgal", command=lambda: self.app.show("EmployeeSelectPage")).pack(side="left", padx=10)
        self.next_btn = ttk.Button(bf, text="Toliau", command=self.on_next)
        self.next_btn.pack(side="left", padx=10)

    def set_today(self):
        self.issue_date = date.today()
        self.issue_date_var.set(self.issue_date.isoformat())

    def pick_date(self):
        init = parse_date(self.issue_date_var.get()) or date.today()

        def on_pick(d: date):
            self.issue_date = d
            self.issue_date_var.set(d.isoformat())

        pop = CalendarPopup(self, init, on_pick)
        pop.geometry(f"+{self.winfo_rootx()+160}+{self.winfo_rooty()+140}")

    def open_dept_list(self):
        values = self.app.list_departments()
        if not values:
            messagebox.showwarning("Sąrašas", "Darbuotojai.xlsx faile nerasta padalinių.")
            return

        def pick(v):
            self.dept_var.set(v)
            self._on_dept_change()

        pop = ListSelectPopup(self, "Padaliniai", values, pick)
        pop.geometry(f"+{self.winfo_rootx()+140}+{self.winfo_rooty()+120}")

    def open_pos_list(self):
        dept = (self.dept_var.get() or "").strip()
        values = self.app.list_positions_for_department(dept) if dept else []
        if not values:
            messagebox.showwarning("Sąrašas", "Šitam padaliniui pareigų nerasta (Darbuotojai.xlsx).")
            return

        def pick(v):
            self.pos_var.set(v)

        pop = ListSelectPopup(self, "Pareigos", values, pick)
        pop.geometry(f"+{self.winfo_rootx()+160}+{self.winfo_rooty()+140}")

    def _on_dept_change(self):
        dept = (self.dept_var.get() or "").strip()
        positions = self.app.list_positions_for_department(dept) if dept else []
        self.pos_combo["values"] = positions
        if self.pos_var.get() not in positions:
            self.pos_var.set(positions[0] if positions else "")

    def refresh(self):
        depts = self.app.list_departments()
        self.dept_combo["values"] = depts

        if self.app.change_mode:
            self.next_btn.config(text="Išsaugoti")
        else:
            self.next_btn.config(text="Toliau")

        e = self.app.current_employee
        if e:
            self.tab_var.set(e.tab_nr)
            self.name_var.set(e.name)
            self.gender_var.set(e.gender or "Vyras")
            self.dept_var.set(e.department if e.department in depts else (depts[0] if depts else ""))
            self._on_dept_change()
            if e.position and e.position in self.pos_combo["values"]:
                self.pos_var.set(e.position)
        else:
            self.tab_var.set("")
            self.name_var.set("")
            self.gender_var.set("Vyras")
            self.dept_var.set(depts[0] if depts else "")
            self._on_dept_change()
            self.set_today()

        if self.app.change_mode:
            self.date_row.grid_forget()
            self.tab_entry.config(state="disabled")
            self.name_entry.config(state="disabled")
        else:
            self.date_row.grid(row=5, column=1, sticky="w", padx=5, pady=6)
            self.tab_entry.config(state="normal")
            self.name_entry.config(state="normal")

    def on_next(self):
        tab_nr = (self.tab_var.get() or "").strip()
        name = (self.name_var.get() or "").strip()
        dept = (self.dept_var.get() or "").strip()
        pos = (self.pos_var.get() or "").strip()
        gender = (self.gender_var.get() or "").strip() or "Vyras"

        if not tab_nr or not name:
            messagebox.showwarning("Trūksta duomenų", "Reikia Tab. Nr ir Vardas Pavardė.")
            return
        if not dept:
            messagebox.showwarning("Trūksta duomenų", "Pasirink Padalinį.")
            return
        if not pos:
            messagebox.showwarning("Trūksta duomenų", "Pasirink Pareigas.")
            return

        e = EmployeeInfo(
            tab_nr=tab_nr,
            name=name,
            department=dept,
            position=pos,
            gender=gender,
            issue_date=parse_date(self.issue_date_var.get()) or date.today(),
        )

        self.app.upsert_employee(e)
        self.app.current_employee = e

        if self.app.change_mode:
            self.app.change_mode = False
            self.app.show("NewGearPage")
        else:
            self.app.show("CurrentGearPage")


# -----------------------------
# Page 3: current gear
# -----------------------------

class CurrentGearPage(ttk.Frame):
    def __init__(self, parent, app):
        super().__init__(parent)
        self.app = app

        ttk.Label(self, text="3) Dabartinė įranga (tik dabartinis padalinys)", font=("Arial", 16)).pack(pady=10)

        # Frame su Treeview + Scrollbar
        table_frame = ttk.Frame(self)
        table_frame.pack(padx=20, pady=10, fill="both", expand=True)

        self.tree = ttk.Treeview(
            table_frame,
            columns=("Apranga", "Kodas", "Išduota", "Mėn", "Keisti iki", "Likę laikas"),
            show="headings",
            height=15,
        )

        y_scroll = ttk.Scrollbar(table_frame, orient="vertical", command=self.tree.yview)
        self.tree.configure(yscrollcommand=y_scroll.set)

        self.tree.pack(side="left", fill="both", expand=True)
        y_scroll.pack(side="right", fill="y")

        # Header + plotis
        for col in self.tree["columns"]:
            self.tree.heading(col, text=col)

            if col == "Apranga":
                self.tree.column(col, width=360, anchor="w")
            elif col == "Kodas":
                self.tree.column(col, width=120, anchor="w")
            elif col in ("Išduota", "Keisti iki"):
                self.tree.column(col, width=110, anchor="center")
            elif col == "Mėn":
                self.tree.column(col, width=70, anchor="center")
            elif col == "Likę laikas":
                self.tree.column(col, width=95, anchor="center")
            else:
                self.tree.column(col, width=110)

        btn = ttk.Frame(self)
        btn.pack(pady=10)

        ttk.Button(btn, text="Atgal", command=lambda: self.app.show("EmployeeSelectPage")).pack(side="left", padx=8)
        ttk.Button(btn, text="Toliau → (nauja įranga)", command=lambda: self.app.show("NewGearPage")).pack(side="left", padx=8)

        # raudonas tekstas, kai < 7 d.
        self.tree.tag_configure("warn", foreground="red")

    def refresh(self):
        # išvalom lentelę
        for iid in self.tree.get_children():
            self.tree.delete(iid)

        employee = getattr(self.app, "current_employee", None)
        if employee is None:
            return

        # !!! svarbiausia: imam iš self.app, ne iš controller
        rows = self.app.load_employee_rows()

        relevant = [r for r in rows if str(r.get("Tab. Nr", "")).strip() == str(employee.tab_nr).strip()]

        # atnaujinam darbuotojo dabartinį padalinį/pareigas/lytį pagal naujausią įrašą
        latest_info = None
        latest_date = None
        for r in relevant:
            issued = parse_date(r.get("Išduota"))
            if issued and (latest_date is None or issued > latest_date):
                latest_date = issued
                latest_info = r

        if latest_info:
            employee.department = str(latest_info.get("Padalinys", "") or "")
            employee.position = str(latest_info.get("Pareigos", "") or "")
            employee.gender = str(latest_info.get("Lytis", "") or "")

        # rodom tik dabartinio padalinio + pareigų įrangą
        filtered = [
            r for r in relevant
            if str(r.get("Padalinys", "") or "") == str(employee.department)
            and str(r.get("Pareigos", "") or "") == str(employee.position)
        ]

        # grupuojam pagal pavadinimą ignoruojant "(xx dydis)"
        latest_by_item = {}
        for r in filtered:
            name = str(r.get("Apranga", "") or "").strip()
            if not name:
                continue

            base_name = strip_size_suffix(name)
            issued = parse_date(r.get("Išduota"))
            if issued is None:
                continue

            ex = latest_by_item.get(base_name)
            if ex is None or issued > ex["issued"]:
                latest_by_item[base_name] = {"row": r, "issued": issued}

        today = date.today()

        for item in latest_by_item.values():
            r = item["row"]
            issued = parse_date(r.get("Išduota"))

            try:
                months = int(r.get("Susidėvėjimas") or 0)
            except Exception:
                months = 0

            change_date = add_months(issued, months) if issued else None
            remaining = (change_date - today).days if change_date else ""

            values = (
                r.get("Apranga", "") or "",
                r.get("Aprangos kodas", "") or "",
                issued.strftime("%Y-%m-%d") if issued else "",
                str(months) if months else "",
                change_date.strftime("%Y-%m-%d") if change_date else "",
                remaining,
            )

            iid = self.tree.insert("", tk.END, values=values)
            if isinstance(remaining, int) and remaining < 7:
                self.tree.item(iid, tags=("warn",))



# -----------------------------
# Page 4: new gear (Kodas | Susidėvėjimas | Apranga RO)
# -----------------------------
class NewGearPage(ttk.Frame):
    def __init__(self, parent, app: AAPApp):
        super().__init__(parent)
        self.app = app

        ttk.Label(self, text="4) Nauja įranga – max 14 eilučių", font=("Arial", 16)).pack(pady=10)

        header = ttk.Frame(self)
        header.pack(padx=20, fill="x")
        ttk.Label(header, text="Kodas", width=20).grid(row=0, column=0, sticky="w")
        ttk.Label(header, text="Susidėvėjimas (mėn.)", width=22).grid(row=0, column=1, sticky="w")
        ttk.Label(header, text="Apranga", width=60).grid(row=0, column=2, sticky="w")

        grid = ttk.Frame(self)
        grid.pack(padx=20, pady=6, fill="x")

        self.rows = []
        self._after = {}
        self._asked_unknown = set()  # idx kuriems jau klausėm pavadinimo

        for i in range(14):
            v_code = tk.StringVar()
            v_months = tk.StringVar()
            v_name = tk.StringVar()

            e_code = ttk.Entry(grid, textvariable=v_code, width=22)
            e_months = ttk.Entry(grid, textvariable=v_months, width=24)
            e_name = ttk.Entry(grid, textvariable=v_name, width=70, state="readonly")
            e_name.configure(takefocus=0)

            e_code.grid(row=i, column=0, pady=3, sticky="w")
            e_months.grid(row=i, column=1, pady=3, padx=(8, 8), sticky="w")
            e_name.grid(row=i, column=2, pady=3, sticky="we")

            grid.columnconfigure(2, weight=1)

            # auto-fill
            e_code.bind("<KeyRelease>", lambda ev, idx=i: self.on_code_change(idx))

            # Tab: Kodas -> Susidėvėjimas
            e_code.bind("<Tab>", lambda ev, idx=i: self.focus_months(idx, ev))

            # Tab: Susidėvėjimas -> kitos eilutės Kodas
            e_months.bind("<Tab>", lambda ev, idx=i: self.focus_next_code(idx, ev))

            self.rows.append({
                "code": v_code,
                "months": v_months,
                "name": v_name,
                "e_code": e_code,
                "e_months": e_months,
            })

        bf = ttk.Frame(self)
        bf.pack(pady=12)
        ttk.Button(bf, text="Atgal", command=lambda: self.app.show("CurrentGearPage")).pack(side="left", padx=8)
        ttk.Button(bf, text="Generuoti ir naujas įrašas", command=self.on_generate_new).pack(side="left", padx=8)
        ttk.Button(bf, text="Generuoti ir uždaryti", command=self.on_generate_close).pack(side="left", padx=8)

    def refresh(self):
        self._asked_unknown.clear()
        self.app.reload_gear_codes()
        for r in self.rows:
            r["code"].set("")
            r["months"].set("")
            r["name"].set("")

    def focus_months(self, idx, _event):
        # jei kodas nežinomas – paprašom pavadinimo (vieną kartą)
        self.maybe_prompt_unknown_code(idx)
        self.rows[idx]["e_months"].focus_set()
        return "break"

    def focus_next_code(self, idx, _event):
        nxt = idx + 1
        if nxt < len(self.rows):
            self.rows[nxt]["e_code"].focus_set()
        else:
            self.rows[idx]["e_months"].focus_set()
        return "break"

    def on_code_change(self, idx):
        if idx in self._after:
            self.after_cancel(self._after[idx])
        self._after[idx] = self.after(120, lambda: self.apply_code(idx))

    def apply_code(self, idx):
        r = self.rows[idx]
        raw = (r["code"].get() or "").strip()
        if not raw:
            r["name"].set("")
            return

        base, size = split_code_and_size(raw)
        info = self.app.gear_codes.get(base)

        if info:
            nm = str(info.get("name", "") or "")
            r["name"].set(ensure_size_suffix(nm, size))
        else:
            # kol kas tuščia – pavadinimą prašysim, kai user eis TAB į susidėvėjimą
            r["name"].set("")

    def maybe_prompt_unknown_code(self, idx):
        if idx in self._asked_unknown:
            return

        r = self.rows[idx]
        raw = (r["code"].get() or "").strip()
        if not raw:
            return

        base, size = split_code_and_size(raw)
        if base in self.app.gear_codes:
            return

        # nežinomas kodas -> paklausti pavadinimo
        self._asked_unknown.add(idx)

        title = "Naujas kodas"
        prompt = f"Kodas '{base}' nerastas „Aprangos kodai.xlsx“.\nĮvesk prekės pavadinimą:"
        name = simpledialog.askstring(title, prompt, parent=self)

        if name:
            name = name.strip()
            # įrašom į Aprangos kodai.xlsx ir perskaitom atgal
            self.app.ensure_gear_code_exists(base, name)
            # užpildom su dydžiu, jei buvo
            self.rows[idx]["name"].set(ensure_size_suffix(name, size))
        else:
            # paliekam tuščią; tokia eilutė bus praleista generuojant
            self.rows[idx]["name"].set("")

    def collect_items(self) -> list[dict]:
        out = []
        for r in self.rows:
            code_raw = (r["code"].get() or "").strip()
            months_raw = (r["months"].get() or "").strip()
            name = (r["name"].get() or "").strip()

            if not code_raw and not months_raw and not name:
                continue

            base, size = split_code_and_size(code_raw)
            base = base.strip()

            if not name:
                # praleidžiam nepilną eilutę
                continue

            try:
                months = int(months_raw)
            except Exception:
                months = 0

            out.append({
                "code": base,  # be "-dydis"
                "name": ensure_size_suffix(strip_size_suffix(name), size),
                "months": months,
            })
        return out

    def run_generation(self, close_after: bool):
        e = self.app.current_employee
        if not e:
            messagebox.showwarning("Klaida", "Nėra pasirinkto darbuotojo.")
            return

        items = self.collect_items()
        if not items:
            messagebox.showwarning("Trūksta duomenų", "Įvesk bent vieną eilutę: Kodas + Susidėvėjimas.")
            return

        # data: naujam darbuotojui gali būti parinkta, esamam – visada šiandien
        issued_date = date.today()
        if e.issue_date and (self.app.change_mode is False):
            issued_date = e.issue_date

        doc_no = self.app.repo_db.next_document_number()

        self.app.repo_db.append_issuance(e, items, doc_no, issued_date)
        self.app.generate_word_doc(e, items, doc_no, issued_date)

        if close_after:
            self.app.destroy()
        else:
            self.app.current_employee = None
            self.app.show("EmployeeSelectPage")

    def on_generate_close(self):
        # be thread – stabiliau, mažiau Tk klaidų
        self.run_generation(True)

    def on_generate_new(self):
        self.run_generation(False)


# -----------------------------
# Start
# -----------------------------
if __name__ == "__main__":
    settings_file = os.environ.get("AAP_SETTINGS", "settings.json")
    app = AAPApp(settings_file)
    app.mainloop()
